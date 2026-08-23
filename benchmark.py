import time
import logging
import asyncio
import httpx
from pathlib import Path
from docx import Document

# Configure logging to show telemetry
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("telemetry")

BASE_URL = "http://localhost:8000"

def create_temp_resume() -> str:
    """Generate a realistic test resume DOCX for benchmarking."""
    doc = Document()
    doc.add_paragraph("Anubhav Sharma")
    doc.add_paragraph("anubhav@example.com | +91-9999999999")
    doc.add_paragraph("Skills")
    doc.add_paragraph("Python, FastAPI, LlamaIndex, Qdrant, Redis, Docker, PyTorch, RAG")
    doc.add_paragraph("Experience")
    doc.add_paragraph("Senior GenAI Engineer at Wayfarer. Built a high-performance RAG pipeline.")
    doc.add_paragraph("Implemented Redis caching and Qdrant integration to reduce response times.")
    doc.add_paragraph("Education")
    doc.add_paragraph("B.Tech CS, IIT Delhi")
    
    file_path = "benchmark_test_resume.docx"
    doc.save(file_path)
    return file_path

async def run_benchmark():
    file_path = create_temp_resume()
    
    async with httpx.AsyncClient(timeout=600) as client:
        print(f"{'Endpoint':<25} | {'Cold (s)':<10} | {'Warm (s)':<10} | {'Improvement'}")
        print("-" * 65)

        # 1. Search Benchmark
        query = {"query": "how to build a RAG system?", "max_sources": 3}
        start = time.perf_counter()
        r1 = await client.post(f"{BASE_URL}/api/v1/search", json=query)
        c1 = time.perf_counter() - start
        
        start = time.perf_counter()
        r2 = await client.post(f"{BASE_URL}/api/v1/search", json=query)
        w1 = time.perf_counter() - start
        
        print(f"{'/api/v1/search':<25} | {c1:<10.4f} | {w1:<10.4f} | {((c1-w1)/c1)*100:.1f}%")

        # 2. Resume Check Benchmark (real parsing + keyword extraction)
        jd_text = "Looking for a Python developer with FastAPI, Qdrant, and Redis experience to build LLM pipelines."
        
        with open(file_path, "rb") as f:
            files = {"resume_file": (file_path, f.read(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        
        start = time.perf_counter()
        r3 = await client.post(
            f"{BASE_URL}/api/v1/resume/check",
            files=files,
            data={"jd_text": jd_text}
        )
        c2 = time.perf_counter() - start
        data3 = r3.json()
        resume_id = data3.get("resume_id")
        
        # Warm check (same file upload & same JD)
        with open(file_path, "rb") as f:
            files = {"resume_file": (file_path, f.read(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        start = time.perf_counter()
        r4 = await client.post(
            f"{BASE_URL}/api/v1/resume/check",
            files=files,
            data={"jd_text": jd_text}
        )
        w2 = time.perf_counter() - start
        
        print(f"{'/api/v1/resume/check':<25} | {c2:<10.4f} | {w2:<10.4f} | {((c2-w2)/c2)*100:.1f}%")

        # 3. Jobs Match Benchmark (RAG Pipeline via LlamaIndex + Qdrant)
        if resume_id:
            params = {"resume_id": resume_id}
            start = time.perf_counter()
            r5 = await client.get(f"{BASE_URL}/api/v1/jobs/match", params=params)
            c3 = time.perf_counter() - start
            
            start = time.perf_counter()
            r6 = await client.get(f"{BASE_URL}/api/v1/jobs/match", params=params)
            w3 = time.perf_counter() - start
            
            print(f"{'/api/v1/jobs/match':<25} | {c3:<10.4f} | {w3:<10.4f} | {((c3-w3)/c3)*100:.1f}%")
        else:
            print(f"{'/api/v1/jobs/match':<25} | {'SKIP':<10} | {'SKIP':<10} | N/A")

    # Clean up temp file
    try:
        Path(file_path).unlink()
    except OSError:
        pass

if __name__ == "__main__":
    asyncio.run(run_benchmark())
