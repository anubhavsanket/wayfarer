"""End-to-end integration tests that run against the live Docker stack.

These tests require all Docker services to be running (api, chromadb, ollama,
redis). They exercise the full API surface against real services.

Run: pytest tests/test_docker_e2e.py -v --timeout=300

The tests are ordered so earlier tests don't depend on later ones:
1. Health check
2. Search (requires NVIDIA NIM + Tavily)
3. Resume check (requires NVIDIA NIM for keyword extraction)
4. Resume save
5. Job match (requires bluedoor API key)
6. Background refresh
"""
from __future__ import annotations

import os
import sys
from pathlib import Path

import httpx
import pytest

# Skip if not in Docker environment
DOCKER_HOST = os.environ.get("DOCKER_HOST", "http://localhost:8000")
pytestmark = pytest.mark.skipif(
    os.environ.get("SKIP_DOCKER_TESTS", "0") == "1",
    reason="SKIP_DOCKER_TESTS=1",
)

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def api(method: str, path: str, **kwargs) -> httpx.Response:
    """Make a request against the API. Returns the response object."""
    url = f"{DOCKER_HOST}{path}"
    with httpx.Client(timeout=300) as client:
        return getattr(client, method)(url, **kwargs)


def create_test_resume(tmp_path: Path = None) -> Path:
    """Create a minimal DOCX resume for testing."""
    import tempfile
    tmp_path = tmp_path or Path(tempfile.gettempdir())
    from docx import Document
    doc = Document()
    doc.add_paragraph("Test User")
    doc.add_paragraph("test@example.com | +91-9999999999")
    doc.add_paragraph("Summary")
    doc.add_paragraph("Engineer focused on RAG systems and LLM applications.")
    doc.add_paragraph("Skills")
    doc.add_paragraph("Python, FastAPI, ChromaDB, LangChain, AWS, PyTorch")
    doc.add_paragraph("Experience")
    doc.add_paragraph("- Built a RAG system with ChromaDB and Ollama embeddings.")
    doc.add_paragraph("- Shipped a FastAPI service serving 10k requests/day.")
    doc.add_paragraph("Education")
    doc.add_paragraph("B.Tech CS, IIT Delhi")
    path = tmp_path / "test_resume.docx"
    doc.save(str(path))
    return path


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------

class TestHealth:
    def test_health_endpoint_returns_200(self):
        resp = api("get", "/health")
        assert resp.status_code == 200

    def test_health_shows_dependencies(self):
        data = api("get", "/health").json()
        assert "status" in data
        assert "dependencies" in data
        deps = {d["name"] for d in data["dependencies"]}
        assert "chromadb" in deps
        assert "ollama" in deps

    def test_health_chromadb_is_up(self):
        data = api("get", "/health").json()
        chromadb = next(d for d in data["dependencies"] if d["name"] == "chromadb")
        assert chromadb["status"] == "up"


class TestSearch:
    def test_search_returns_answer(self):
        resp = api("post", "/api/v1/search", json={"query": "what is RAG?", "max_sources": 2})
        assert resp.status_code == 200
        data = resp.json()
        assert "answer" in data
        assert "citations" in data
        assert "sub_queries_used" in data
        assert len(data["answer"]) > 0

    def test_search_has_citations(self):
        resp = api("post", "/api/v1/search", json={"query": "what is RAG?", "max_sources": 2})
        data = resp.json()
        # Citations should be a list (may be empty if synthesis failed)
        assert isinstance(data["citations"], list)

    def test_search_invalid_query_rejected(self):
        resp = api("post", "/api/v1/search", json={"query": "", "max_sources": 2})
        assert resp.status_code == 422


class TestResumeCheck:
    def test_resume_check_with_docx(self):
        resume_path = create_test_resume()
        with open(resume_path, "rb") as f:
            resp = api(
                "post",
                "/api/v1/resume/check",
                files={"resume_file": ("test.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                data={"jd_text": "Looking for an ML engineer with Python, PyTorch, FastAPI, and AWS experience."},
            )
        assert resp.status_code == 200
        data = resp.json()
        assert "resume_id" in data
        assert "ats_score" in data
        assert "keyword_gaps" in data
        assert data["ats_score"] >= 0.0
        assert len(data["keyword_gaps"]) > 0
        # Verify confidence tiers are valid
        for gap in data["keyword_gaps"]:
            assert gap["tier"] in ("verified", "reworded", "gap")

    def test_resume_check_missing_file_rejected(self):
        resp = api("post", "/api/v1/resume/check", data={"jd_text": "test"})
        assert resp.status_code == 422


class TestResumeSave:
    def test_save_creates_new_file(self):
        # First check a resume to get a valid resume_id
        resume_path = create_test_resume()
        with open(resume_path, "rb") as f:
            resp = api(
                "post",
                "/api/v1/resume/check",
                files={"resume_file": ("test.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                data={"jd_text": "Python engineer"},
            )
        resume_id = resp.json()["resume_id"]
        # Save with accepted suggestions
        resp = api("post", "/api/v1/resume/save", json={
            "resume_id": resume_id,
            "accepted_suggestions": [{"bullet_id": "b1", "suggested_text": "Built RAG with ChromaDB."}],
            "mode": "new_file",
            "confirm_overwrite": False,
        })
        assert resp.status_code == 200
        data = resp.json()
        assert "file_id" in data
        assert data["mode_applied"] == "new_file"


class TestJobMatch:
    def test_job_match_with_mock(self):
        resp = api("get", "/api/v1/jobs/match", params={"resume_id": "test", "limit": 3, "test": "true"})
        assert resp.status_code == 200
        data = resp.json()
        assert "matches" in data
        assert "aggregate_gaps" in data

    def test_job_match_invalid_resume_id(self):
        resp = api("get", "/api/v1/jobs/match", params={"resume_id": "invalid"})
        assert resp.status_code == 400


class TestBackgroundRefresh:
    def test_refresh_returns_counts(self):
        resp = api("post", "/api/v1/jobs/refresh")
        assert resp.status_code == 200
        data = resp.json()
        assert "refreshed" in data
        assert data["refreshed"] > 0
        assert "by_source" in data
