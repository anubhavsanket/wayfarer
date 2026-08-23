"""Unit tests for core services (parser, ats_checker, job_matcher).

This expands test coverage beyond integration/E2E tests.
"""
import os
import sys
import tempfile
import pytest
from pathlib import Path
from docx import Document

BACKEND = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(BACKEND))

# Ensure no API keys for tests
for key in ("TAVILY_API_KEY", "BRAVE_API_KEY", "NVIDIA_NIM_API_KEY", "OPENROUTER_API_KEY"):
    os.environ.pop(key, None)


# ---------------------------------------------------------------------------
# Resume Parser
# ---------------------------------------------------------------------------

class TestResumeParser:
    def test_parser_extracts_sections(self):
        from backend.app.services.resume_parser import parse_resume
        doc = Document()
        doc.add_paragraph("Test User")
        doc.add_paragraph("test@example.com")
        doc.add_paragraph("Skills")
        doc.add_paragraph("Python, FastAPI")
        doc.add_paragraph("Experience")
        doc.add_paragraph("Built RAG systems")
        
        path = os.path.join(tempfile.gettempdir(), "test_parser_sections.docx")
        doc.save(path)
        
        parsed = parse_resume(path)
        assert "skills" in parsed.sections
        assert "experience" in parsed.sections
        assert len(parsed.bullets) > 0

    def test_parser_extracts_bullets(self):
        from backend.app.services.resume_parser import parse_resume
        doc = Document()
        doc.add_paragraph("Test User")
        doc.add_paragraph("Experience")
        doc.add_paragraph("Built RAG systems with ChromaDB and FastAPI")
        doc.add_paragraph("Led a team of 3")
        
        path = os.path.join(tempfile.gettempdir(), "test_parser_bullets.docx")
        doc.save(path)
        
        parsed = parse_resume(path)
        assert len(parsed.bullets) == 2
        assert "ChromaDB" in parsed.bullets[0].text

    def test_parser_contact_extraction(self):
        from backend.app.services.resume_parser import parse_resume
        doc = Document()
        doc.add_paragraph("Test User")
        doc.add_paragraph("test@example.com | +91-9999999999")
        doc.add_paragraph("Summary")
        doc.add_paragraph("Engineer with experience.")
        
        path = os.path.join(tempfile.gettempdir(), "test_parser_contact.docx")
        doc.save(path)
        
        parsed = parse_resume(path)
        assert parsed.contact.get("email") == "test@example.com"
        assert parsed.contact.get("phone") == "+91-9999999999"


# ---------------------------------------------------------------------------
# ATS Checker
# ---------------------------------------------------------------------------

class TestATSChecker:
    def test_fallback_keywords(self):
        from backend.app.services.ats_checker import _fallback_keywords
        jd = "Looking for Python developer with FastAPI and Docker experience"
        keywords = _fallback_keywords(jd)
        assert "python" in keywords
        assert "fastapi" in keywords
        assert "docker" in keywords

    def test_fallback_keywords_max_limit(self):
        from backend.app.services.ats_checker import _fallback_keywords
        jd = "Python Java C++ Rust Go TypeScript JavaScript React Angular Vue Django Flask FastAPI"
        keywords = _fallback_keywords(jd)
        assert len(keywords) <= 25


# ---------------------------------------------------------------------------
# Job Matcher
# ---------------------------------------------------------------------------

class TestJobMatcher:
    def _make_job(self, title: str = "Job1", gaps: list[str] | None = None):
        from backend.app.models.schemas import JobMatch, LocationMatch
        return JobMatch(
            job_id="j1",
            title=title,
            company="C1",
            source="test",
            location="Remote",
            match_score=0.8,
            location_match=LocationMatch.NONE,
            top_gaps=gaps or [],
            apply_url="https://example.com",
        )

    def test_aggregate_gaps(self):
        from backend.app.services.job_matcher import _aggregate_gaps
        
        matches = [
            self._make_job("Job1", ["Rust", "Go"]),
            self._make_job("Job2", ["Rust", "Python"]),
            self._make_job("Job3", ["Go"]),
        ]
        
        gaps = _aggregate_gaps(matches)
        assert len(gaps) > 0
        assert gaps[0].skill == "Rust"  # Most common gap

    def test_aggregate_gaps_empty(self):
        from backend.app.services.job_matcher import _aggregate_gaps
        
        matches = [
            self._make_job("Job1", []),
        ]
        
        gaps = _aggregate_gaps(matches)
        assert len(gaps) == 0
