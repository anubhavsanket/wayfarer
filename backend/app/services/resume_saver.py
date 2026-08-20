"""Resume save service — apply accepted suggestions, write a new file.

FR2.8 (PRD §8.1): After the user reviews and accepts/rejects individual
redline suggestions, let them choose to **save as a new file** (default,
non-destructive) or **overwrite the original** upload. Overwrite requires
an explicit ``confirm_overwrite`` — never silently replace the source.

v2: OOXML track-changes support via <w:ins> and <w:del> elements so the
saved .docx opens in Word with changes marked for interactive accept/reject.
"""
from __future__ import annotations

import logging
import re
import uuid
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from ..models.schemas import (
    AcceptedSuggestion,
    ChangeSummary,
    ResumeSaveResponse,
    SaveMode,
)
from .resume_parser import ParsedResume
from . import resume_store

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# OOXML track-changes helpers
# ---------------------------------------------------------------------------

def _make_track_change_elements(
    original_text: str,
    suggested_text: str,
    author: str,
    change_id: int,
) -> tuple[OxmlElement, OxmlElement]:
    """Create <w:del> and <w:ins> elements for a single change."""
    now = datetime.now(timezone.utc).isoformat().replace("+00:00", "Z")

    # Deletion element (original text)
    del_el = OxmlElement("w:del")
    del_el.set(qn("w:id"), str(change_id))
    del_el.set(qn("w:author"), author)
    del_date = OxmlElement("w:date")
    del_date.text = now
    del_el.append(del_date)

    r_del = OxmlElement("w:r")
    t_del = OxmlElement("w:delText")
    t_del.text = original_text
    r_del.append(t_del)
    del_el.append(r_del)

    # Insertion element (suggested text)
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), str(change_id + 1))
    ins.set(qn("w:author"), author)
    ins_date = OxmlElement("w:date")
    ins_date.text = now
    ins.append(ins_date)

    r_ins = OxmlElement("w:r")
    t_ins = OxmlElement("w:t")
    t_ins.text = suggested_text
    r_ins.append(t_ins)
    ins.append(r_ins)

    return del_el, ins


# ---------------------------------------------------------------------------
# DOCX writer with track-changes
# ---------------------------------------------------------------------------

def _write_docx(
    parsed: ParsedResume,
    out_path: Path,
    accepted_suggestions: list[AcceptedSuggestion] | None = None,
) -> dict[str, int]:
    """Reconstruct a .docx from parsed sections with OOXML track-changes.

    Args:
        parsed: The parsed resume (with edits already applied to bullets).
        out_path: Output file path.
        accepted_suggestions: The original accepted suggestions, used to generate
            track-changes. If None, writes a clean document without track-changes.

    Returns:
        Dict with counts: {"insertions": n, "deletions": n, "total_changes": n}
    """
    doc = Document()

    # Build a map of bullet_id -> (original_text, suggested_text) for track-changes
    track_changes_map: dict[str, tuple[str, str]] = {}
    if accepted_suggestions:
        for s in accepted_suggestions:
            if s.original_text and s.suggested_text:
                track_changes_map[s.bullet_id] = (s.original_text, s.suggested_text)

    change_id = 0
    counts = {"insertions": 0, "deletions": 0, "total_changes": 0}

    # Title / name
    name = parsed.contact.get("name") or "Resume"
    doc.add_heading(name, level=0)

    # Contact line
    contact_parts = [
        parsed.contact.get("email", ""),
        parsed.contact.get("phone", ""),
        parsed.contact.get("linkedin", ""),
        parsed.contact.get("github", ""),
    ]
    contact_line = " | ".join(p for p in contact_parts if p)
    if contact_line:
        doc.add_paragraph(contact_line)

    # Sections
    section_titles = {
        "summary": "Summary",
        "skills": "Skills",
        "experience": "Experience",
        "projects": "Projects",
        "education": "Education",
        "certifications": "Certifications",
        "awards": "Awards",
        "publications": "Publications",
    }

    # Map bullet id → final text (already applied in _apply_suggestions)
    bullet_final_text: dict[str, str] = {b.id: b.text for b in parsed.bullets}

    for section_name, lines in parsed.sections.items():
        if section_name in ("contact", "preamble"):
            continue
        title = section_titles.get(section_name, section_name.replace("_", " ").title())
        doc.add_heading(title, level=1)
        for line in lines:
            # Find the bullet this line corresponds to
            bullet_match = next(
                (b for b in parsed.bullets if b.text == line),
                None,
            )
            is_bullet = line.strip().startswith(("•", "-", "*"))
            style = "List Bullet" if is_bullet else None

            if bullet_match and bullet_match.id in track_changes_map:
                # This line has a track-change: write both deletion and insertion
                original, suggested = track_changes_map[bullet_match.id]
                p = doc.add_paragraph(style=style)
                p_xml = p._p

                # Clear any default runs
                for run in p.runs[:]:
                    run._element.getparent().remove(run._element)

                # Write deletion (original)
                del_el, ins_el = _make_track_change_elements(
                    original, suggested, "Wayfarer", change_id
                )
                p_xml.append(del_el)
                p_xml.append(ins_el)

                change_id += 2
                counts["deletions"] += 1
                counts["insertions"] += 1
                counts["total_changes"] += 1
            else:
                # No change - write the current text
                text = bullet_final_text.get(bullet_match.id, line) if bullet_match else line
                doc.add_paragraph(text, style=style)

    doc.save(str(out_path))
    return counts


# ---------------------------------------------------------------------------
# Apply suggestions
# ---------------------------------------------------------------------------

_BULLET_GLYPH_RE = re.compile(r"^([•‣◦⁃\-\*]\s*)")


def _strip_bullet_glyph(text: str) -> str:
    """Remove a leading bullet glyph (• - *) from a section line."""
    return _BULLET_GLYPH_RE.sub("", text).strip()


def _with_bullet_glyph(original: str, text: str) -> str:
    """Prepend the glyph used by ``original`` (if any) onto ``text``.

    Keeps a List-Bullet paragraph looking like a bullet after an edit.
    """
    match = _BULLET_GLYPH_RE.match(original)
    return f"{match.group(1)}{text}" if match else text


def _apply_suggestions(
    parsed: ParsedResume,
    accepted: list[AcceptedSuggestion],
) -> ParsedResume:
    """Return a copy of ``parsed`` with accepted bullet edits applied."""
    edits = {s.bullet_id: s.suggested_text for s in accepted}

    # Snapshot the ORIGINAL bullet text before applying edits. We match
    # section lines against this snapshot, not the post-edit text, so the
    # section stays consistent with the bullets.
    original_text_by_id = {b.id: b.text for b in parsed.bullets}

    # Update bullets
    for b in parsed.bullets:
        if b.id in edits and edits[b.id]:
            b.text = edits[b.id]

    # Update the section lines that came from edited bullets
    new_sections: dict[str, list[str]] = {}
    for section, lines in parsed.sections.items():
        new_lines = []
        for line in lines:
            clean_line = _strip_bullet_glyph(line)
            match = next(
                (b for b in parsed.bullets if original_text_by_id.get(b.id) == clean_line),
                None,
            )
            if match and match.id in edits:
                new_lines.append(_with_bullet_glyph(line, edits[match.id]))
            else:
                new_lines.append(line)
        new_sections[section] = new_lines
    parsed.sections = new_sections

    return parsed


# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------

async def save_resume(
    resume_id: str,
    accepted_suggestions: list[AcceptedSuggestion],
    mode: SaveMode,
    confirm_overwrite: bool = False,
) -> ResumeSaveResponse:
    """Apply accepted suggestions and save as a new file or overwrite."""
    parsed = resume_store.load_parsed(resume_id)
    if parsed is None:
        raise ValueError(f"Unknown resume_id: {resume_id}. Run /resume/check first.")

    if mode == SaveMode.OVERWRITE and not confirm_overwrite:
        raise ValueError(
            "confirm_overwrite must be true to overwrite the original resume."
        )

    # Apply accepted suggestions to the parsed resume
    edited = _apply_suggestions(parsed, accepted_suggestions)

    # Determine output filename
    original = resume_store.original_file_path(resume_id)
    if mode == SaveMode.OVERWRITE and original is not None:
        out_path = original  # overwrite in place
    else:
        rdir = resume_store._resume_dir(resume_id)
        rdir.mkdir(parents=True, exist_ok=True)
        out_path = rdir / f"saved_{uuid.uuid4().hex[:8]}.docx"

    # Write the DOCX with track-changes
    counts = _write_docx(edited, out_path, accepted_suggestions)

    if mode == SaveMode.OVERWRITE:
        # Re-persist the edited parsed data so future saves reflect the change
        resume_store.save_parsed(resume_id, edited)

    return ResumeSaveResponse(
        file_id=uuid.uuid4().hex[:12],
        file_ref=str(out_path),
        mode_applied=mode,
        changes={
            "total_changes": counts["total_changes"],
            "insertions": counts["insertions"],
            "deletions": counts["deletions"],
            "accepted_count": len([s for s in accepted_suggestions if s.suggested_text]),
            "rejected_count": len([s for s in accepted_suggestions if not s.suggested_text]),
        },
    )
