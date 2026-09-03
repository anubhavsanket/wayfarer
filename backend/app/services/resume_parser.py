"""Resume parser: PDF/DOCX → structured sections + ATS-visible text.

Uses high-speed Markdown parsing first (PyMuPDF4LLM for PDF, MarkItDown for DOCX):
- Converts documents to structured Markdown in <300ms
- Preserves layout, section headers, bullet lists, and table indicators
- Generates layout-blind plain text simulating legacy ATS parsers

Falls back to Unstructured.io or pdfplumber + python-docx when Markdown parsers
are unavailable.

FR2.1 (PRD §8.1): Parse an uploaded resume into structured sections
(contact, skills, experience bullets, education).

FR2.2 (PRD §8.2): Run an ATS parsing simulation — extract text the way a
naive ATS parser would (layout-blind, table-ignorant), and diff against
the properly structured extraction to flag structural loss (tables,
multi-column sections that vanish).

Acceptance criterion: "A resume with a table-based skills section shows a
structural-loss flag on that section."
"""
from __future__ import annotations

import re
import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from ..models.schemas import StructuralIssue
from ..core.confidence import ResumeBullet

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Common section header patterns
# ---------------------------------------------------------------------------

SECTION_PATTERNS = [
    (re.compile(r"^(contact|personal|address|info)\s*$", re.I), "contact"),
    (re.compile(r"^(skills?|technical|technologies|competenc)\w*\s*$", re.I), "skills"),
    (re.compile(r"^(experience|work\s+experience|employment|professional)\s*$", re.I), "experience"),
    (re.compile(r"^(education|academics?|qualifications?)\s*$", re.I), "education"),
    (re.compile(r"^(projects?|portfolio|work)\s*$", re.I), "projects"),
    (re.compile(r"^(certifications?|certificates?|licenses?)\s*$", re.I), "certifications"),
    (re.compile(r"^(summary|profile|objective|about)\s*$", re.I), "summary"),
    (re.compile(r"^(publications?|research|papers?)\s*$", re.I), "publications"),
    (re.compile(r"^(awards?|honors?|achievements?)\s*$", re.I), "awards"),
]

BULLET_PATTERN = re.compile(r"^[•‣◦⁃•\-\*]\s*")
EMAIL_RE = re.compile(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}")
PHONE_RE = re.compile(r"[\+]?[\d\-\(\)\s]{7,15}")


def _detect_section(line: str) -> str | None:
    stripped = line.strip()
    if not stripped:
        return None
    for pattern, name in SECTION_PATTERNS:
        if pattern.match(stripped):
            return name
    return None


def _extract_contact(lines: list[str]) -> dict[str, str]:
    contact: dict[str, str] = {}
    blob = " ".join(lines[:5])
    email = EMAIL_RE.search(blob)
    if email:
        contact["email"] = email.group()
    phone = PHONE_RE.search(blob)
    if phone:
        contact["phone"] = phone.group().strip()
    if lines:
        contact["name"] = lines[0].strip()
    return contact


BULLET_SECTIONS = {"experience", "projects", "awards", "publications", "skills"}


def _parse_sections(lines: list[str]) -> tuple[dict[str, list[str]], list[ResumeBullet]]:
    """Extract named sections and bullets from lines."""
    sections: dict[str, list[str]] = {}
    bullets: list[ResumeBullet] = []
    current_section = "preamble"
    bullet_id = 0

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        detected = _detect_section(stripped)
        if detected:
            current_section = detected
            sections.setdefault(current_section, [])
            continue

        sections.setdefault(current_section, []).append(stripped)
        clean = BULLET_PATTERN.sub("", stripped).strip()
        is_bullet = bool(BULLET_PATTERN.match(stripped)) or (
            current_section in BULLET_SECTIONS and clean
        )
        if is_bullet and clean:
            bullets.append(ResumeBullet(
                id=f"b{bullet_id}",
                section=current_section,
                text=clean,
            ))
            bullet_id += 1

    return sections, bullets



# ---------------------------------------------------------------------------
# High-Speed Markdown Parsing (PyMuPDF4LLM for PDF, MarkItDown for DOCX)
# ---------------------------------------------------------------------------

def _clean_md_header(line: str) -> str:
    """Strip markdown headers (#, ##), bold (**), and underline formatting."""
    s = re.sub(r"^#+\s*", "", line).strip()
    s = re.sub(r"^\*\*(.*?)\*\*$", r"\1", s).strip()
    s = re.sub(r"^__(.*?)__$", r"\1", s).strip()
    return s.rstrip(":")


def _try_markdown_parser(file_path: str) -> ParsedResume | None:
    """Attempt high-speed Markdown parsing. Returns None on failure."""
    path = Path(file_path)
    suffix = path.suffix.lower()
    md_text: str | None = None

    if suffix == ".pdf":
        try:
            import pymupdf4llm
            md_text = pymupdf4llm.to_markdown(file_path)
        except Exception as exc:
            logger.debug("PyMuPDF4LLM conversion failed: %s", exc)
            return None
    elif suffix in (".docx", ".doc"):
        try:
            from markitdown import MarkItDown
            converter = MarkItDown()
            res = converter.convert(file_path)
            md_text = res.text_content
        except Exception as exc:
            logger.debug("MarkItDown conversion failed: %s", exc)
            return None
    else:
        return None

    if not md_text or not md_text.strip():
        return None

    sections: dict[str, list[str]] = {}
    current_section = "preamble"
    bullets: list[ResumeBullet] = []
    bullet_id = 0
    all_text_lines: list[str] = []
    structural_issues: list[StructuralIssue] = []

    lines = md_text.splitlines()

    # Detect tables in Markdown
    table_lines = [l for l in lines if "|" in l and ("---" in l or l.strip().startswith("|"))]
    if table_lines:
        structural_issues.append(StructuralIssue(
            location="Global",
            issue="Document contains table structures. A naive ATS may lose formatting or merge cell contents.",
        ))

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # Skip markdown table divider lines
        if re.match(r"^\|?[\s\-:]+(\|[\s\-:]+)+\|?$", stripped):
            continue

        clean_header = _clean_md_header(stripped)
        detected = _detect_section(clean_header) or (
            _detect_section(stripped) if stripped != clean_header else None
        )

        if detected:
            current_section = detected
            sections.setdefault(current_section, [])
            continue

        # Clean text line (strip markdown formatting for content lines)
        clean_line = re.sub(r"^\s*\|?\s*", "", stripped)
        clean_line = re.sub(r"\s*\|\s*$", "", clean_line)
        clean_line = clean_line.replace("|", " ").strip()

        sections.setdefault(current_section, []).append(clean_line)
        all_text_lines.append(clean_line)

        # Extract bullets
        clean_bullet = BULLET_PATTERN.sub("", clean_line).strip()
        clean_bullet = re.sub(r"^\*\*(.*?)\*\*", r"\1", clean_bullet).strip()
        is_bullet = bool(BULLET_PATTERN.match(stripped)) or (
            current_section in BULLET_SECTIONS and clean_bullet
        )
        if is_bullet and clean_bullet:
            bullets.append(ResumeBullet(id=f"b{bullet_id}", section=current_section, text=clean_bullet))
            bullet_id += 1

    contact = _extract_contact(all_text_lines)

    # ATS simulation text (strip markdown symbols for layout-blind ATS view)
    ats_visible_text = re.sub(r"[#\*_`~\[\]]", "", md_text)
    ats_visible_text = re.sub(r"\|", " ", ats_visible_text)
    ats_visible_text = "\n".join(l.strip() for l in ats_visible_text.splitlines() if l.strip())

    raw_text = "\n".join(all_text_lines)

    return ParsedResume(
        sections=sections,
        bullets=bullets,
        ats_visible_text=ats_visible_text,
        raw_text=raw_text,
        structural_issues=structural_issues,
        contact=contact,
    )


# ---------------------------------------------------------------------------
# Unstructured.io parsing (fallback path)
# ---------------------------------------------------------------------------

def _try_unstructured(file_path: str) -> ParsedResume | None:
    """Attempt parsing via Unstructured.io. Returns None on failure."""
    try:
        from unstructured.partition.auto import partition
    except ImportError:
        return None

    try:
        # hi_res for structured extraction
        hi_elements = partition(filename=file_path, strategy="hi_res")
        # fast for ATS simulation
        fast_elements = partition(filename=file_path, strategy="fast")
    except Exception as exc:
        logger.debug("Unstructured partition failed: %s", exc)
        return None

    # ── Build sections from hi_res elements ─────────────────────────────
    sections: dict[str, list[str]] = {}
    current_section = "Contact"
    structural_issues: list[StructuralIssue] = []
    all_text_lines: list[str] = []
    has_table = False
    bullet_id = 0
    bullets: list[ResumeBullet] = []

    for el in hi_elements:
        tag = el.category.lower() if hasattr(el, "category") else ""
        text = str(el).strip()
        if not text:
            continue

        if tag == "table":
            has_table = True
            structural_issues.append(StructuralIssue(
                location=current_section,
                issue=f"Table-based content detected in '{current_section}' section. "
                      f"A naive ATS may lose formatting or merge cell contents.",
            ))

        # Detect section headers
        if tag == "title" or _detect_section(text):
            section_name = _detect_section(text) or text.strip().rstrip(":")
            current_section = section_name
            sections.setdefault(current_section, [])
            continue

        sections.setdefault(current_section, []).append(text)
        all_text_lines.append(text)

        # Extract bullets
        clean = BULLET_PATTERN.sub("", text).strip()
        is_bullet = bool(BULLET_PATTERN.match(text)) or (
            current_section in BULLET_SECTIONS and clean
        )
        if is_bullet and clean:
            bullets.append(ResumeBullet(id=f"b{bullet_id}", section=current_section, text=clean))
            bullet_id += 1

    if has_table:
        structural_issues.append(StructuralIssue(
            location="Global",
            issue="Document contains tables. Verify that table content "
                  "renders correctly in plain-text ATS output.",
        ))

    # ATS simulation text
    ats_text = "\n".join(str(el).strip() for el in fast_elements if str(el).strip())

    # Contact extraction
    contact = _extract_contact(all_text_lines)

    flat_sections = {k: v for k, v in sections.items()}
    raw_text = "\n".join(all_text_lines)

    return ParsedResume(
        sections=flat_sections,
        bullets=bullets,
        ats_visible_text=ats_text,
        raw_text=raw_text,
        structural_issues=structural_issues,
        contact=contact,
    )


# ---------------------------------------------------------------------------
# Fallback parsing (pdfplumber + python-docx)
# ---------------------------------------------------------------------------

def _parse_pdf_fallback(file_path: str) -> tuple[list[str], dict[int, str], list[dict[str, Any]]]:
    import pdfplumber
    all_lines: list[str] = []
    page_texts: dict[int, str] = {}
    tables: list[dict[str, Any]] = []

    with pdfplumber.open(file_path) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            page_text = page.extract_text(x_tolerance=2, y_tolerance=2) or ""
            page_texts[page_num] = page_text
            all_lines.extend(page_text.splitlines())
            for table in page.extract_tables():
                if not table:
                    continue
                headers = [str(c or "").strip() for c in table[0]]
                rows = [[str(c or "").strip() for c in row] for row in table[1:]]
                tables.append({
                    "page": page_num,
                    "headers": headers,
                    "rows": rows,
                    "cell_text": " ".join(" ".join(row) for row in [headers] + rows),
                })

    return all_lines, page_texts, tables


def _parse_docx_fallback(file_path: str) -> tuple[list[str], dict[int, str], list[dict[str, Any]]]:
    import docx as python_docx
    doc = python_docx.Document(file_path)
    paragraph_lines: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraph_lines.append(text)

    tables: list[dict[str, Any]] = []
    for table in doc.tables:
        if not table.rows:
            continue
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows[1:]]
        cell_text = " ".join(" ".join(row) for row in [headers] + rows)
        tables.append({"page": 1, "headers": headers, "rows": rows, "cell_text": cell_text})

    return paragraph_lines, {1: "\n".join(paragraph_lines)}, tables


def _detect_structural_issues_fallback(
    tables: list[dict[str, Any]],
    ats_text: str,
) -> list[StructuralIssue]:
    issues: list[StructuralIssue] = []
    ats_lower = ats_text.lower()
    for tbl in tables:
        for header in tbl["headers"]:
            header_stripped = header.strip().lower()
            if header_stripped and header_stripped not in ats_lower:
                issues.append(StructuralIssue(
                    location=f"table on page {tbl['page']}",
                    issue=(
                        f"Table header '{header}' not visible in ATS text — "
                        f"table structure may be lost by ATS parsers"
                    ),
                ))
    return issues


def _fallback_parse(file_path: str) -> ParsedResume:
    """Parse using pdfplumber/python-docx (legacy path)."""
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext == ".pdf":
        raw_lines, page_texts, tables = _parse_pdf_fallback(str(path))
        ats_text = "\n".join(page_texts.values())
    elif ext in (".docx", ".doc"):
        raw_lines, page_texts, tables = _parse_docx_fallback(str(path))
        ats_text = "\n".join(page_texts.values())
    else:
        raise ValueError(f"Unsupported resume format: {ext}")

    structural_issues = _detect_structural_issues_fallback(tables, ats_text)
    sections, bullets = _parse_sections(raw_lines)
    contact = _extract_contact(raw_lines)
    raw_text = "\n".join(raw_lines)

    return ParsedResume(
        sections=sections,
        bullets=bullets,
        ats_visible_text=ats_text,
        raw_text=raw_text,
        structural_issues=structural_issues,
        contact=contact,
    )


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

@dataclass
class ParsedResume:
    """Structured resume data produced by ``parse_resume``."""
    sections: dict[str, list[str]] = field(default_factory=dict)
    bullets: list[ResumeBullet] = field(default_factory=list)
    ats_visible_text: str = ""
    raw_text: str = ""
    structural_issues: list[StructuralIssue] = field(default_factory=list)
    contact: dict[str, str] = field(default_factory=dict)


def parse_resume(file_path: str, filename: str = "") -> ParsedResume:
    """Parse a resume file (PDF/DOCX) and return structured data + ATS text.

    Tries high-speed Markdown parsing first (PyMuPDF4LLM for PDF, MarkItDown for DOCX).
    Falls back to Unstructured.io, and finally pdfplumber/python-docx.
    """
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix not in (".pdf", ".docx"):
        raise ValueError(f"Unsupported resume format: {suffix}. Only PDF and DOCX are supported.")

    # Try high-speed Markdown parser first
    result = _try_markdown_parser(str(path))
    if result is not None:
        logger.info("Parsed with Markdown parser: %s", file_path)
        return result

    # Try Unstructured.io next
    result = _try_unstructured(str(path))
    if result is not None:
        logger.info("Parsed with Unstructured.io: %s", file_path)
        return result

    # Fallback to pdfplumber/python-docx
    logger.info("Falling back to legacy parser: %s", file_path)
    return _fallback_parse(str(path))
